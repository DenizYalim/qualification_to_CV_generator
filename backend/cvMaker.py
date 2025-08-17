from docx import Document
from docx2pdf import convert
from pathlib import Path
import json

TEMPLATES_FOLDER = Path("../CV_templates")
OUTPUT_FOLDER = Path("../CVs_generated")
OUTPUT_FOLDER.mkdir(parents=True, exist_ok=True)


def _replace_one_placeholder_in_paragraph(p, key: str, val: str) -> bool:
    ph = f"{{{{{key}}}}}"
    if ph not in p.text:
        return False
    txt = p.text
    # replace only the first occurrence for list items
    new_txt = txt.replace(ph, val, 1)
    if new_txt != txt:
        p.text = new_txt  # note: resets runs formatting in this paragraph
        return True
    return False


def _place_value(doc: Document, key: str, val: str) -> bool:
    # paragraphs
    for p in doc.paragraphs:
        if _replace_one_placeholder_in_paragraph(p, key, val):
            return True
    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if _replace_one_placeholder_in_paragraph(p, key, val):
                        return True
    # header/footer
    for sec in doc.sections:
        for part in (sec.header, sec.footer):
            for p in part.paragraphs:
                if _replace_one_placeholder_in_paragraph(p, key, val):
                    return True
    return False


def _replace_in_doc(doc: Document, mapping: dict):
    for k, v in mapping.items():
        key = str(k)
        if isinstance(v, list):
            for item in v:
                _place_value(doc, key, str(item))
        else:
            # replace all occurrences for scalar by looping until none left
            while _place_value(doc, key, str(v)):
                pass


def fill_cv(template_file_name: str, values: dict) -> Path:
    tpl = TEMPLATES_FOLDER / template_file_name
    doc = Document(tpl)
    _replace_in_doc(doc, values)
    out_docx = OUTPUT_FOLDER / f"filled_{Path(template_file_name).stem}.docx"
    doc.save(out_docx)
    return out_docx


# I can overengineer this with regex, or just pass in a json with empty values
# Assumes file is in CVs_generated
def removeUnfilledFields(docName: str):
    docName = OUTPUT_FOLDER / docName
    doc = Document(docName)

    with open("../empty_values.json", "r") as f:
        f = json.load(f)
        _replace_in_doc(doc, f)

    doc.save(docName)


def convertDocToPdf(docPath, pdfPath):
    convert(docPath, pdfPath)


if __name__ == "__main__":
    with open("../example_values.json", "r", encoding="utf-8") as f:
        values = json.load(f)
    template = "cv_gen_template_test.docx"
    out = fill_cv(template, values)
    removeUnfilledFields(out)
    print(f"saved doc to: {out}")
    # pdf = convertDocToPdf(out, "/CVs_generated/a.pdf")
    # print(pdf)
